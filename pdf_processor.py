# coding=utf-8
import fitz  # PyMuPDF
import easyocr
import openai
import docx
from docx import Document
import pandas as pd
import json
import os
import requests
from PDF_OCR import XHD_OCR_Structure
# from win32com import client
# import nltk
# from nltk.tokenize import word_tokenize

from io import BytesIO
import requests
# import office
import chardet

class PDFParser:
    def __init__(self):
        self.model = "gpt-4-0125-preview"
        self.template = self.load_template()
        self.prompt = """你是一个聪明而且有百年经验的命名实体识别（NER）识别器。你的任务是提取字段。你的回答务必使用完全的json格式，json里面请不要嵌套！不要包含多的字段，你要提取的字段有:"""

    def load_template(self):
        # load from file
        template = pd.read_excel("./template.xlsx")
        return template

    def view_template(self):
        # 获取 DataFrame 的列名
        columns = self.template.columns.tolist()
        # 将 DataFrame 中的数据转换为 UTF-8 编码
        records = self.template.applymap(
            lambda x: str(x).encode('utf-8').decode('utf-8') if isinstance(x, str) else x
        ).to_records(index=False)
        # 使用列名和 records 创建字典列表
        template_list = [dict(zip(columns, row)) for row in records]
        # 将数据转换为 JSON，并确保不将中文转换为 Unicode 转义序列
        json_data = json.dumps(template_list, ensure_ascii=False, indent=4)
        return json_data

    def update_template(self, new_template):
        # 格式检查，检查前端传过来的是否符合要求
        self.template = pd.DataFrame(new_template)
        # 返回状态码，是否成功
        return 200

    def view_prompt(self):
        return self.prompt

    def update_prompt(self, prompt):
        self.prompt = prompt
        # 返回状态码，是否成功
        return 200

    def view_model(self):
        return self.model

    def update_model(self, model_name):
        self.model = model_name
        # 返回状态码，是否成功
        return 200

    def pdf_parser(self, template, input_path, model_name = 'gpt-4-0125-preview'):

        # parse the pdf file
        if isinstance(input_path, bytes):
            # 处理字节流
            json_string = {}  # 创建一个空的字典用于存储JSON数据
            # todo
        elif isinstance(input_path, str):
            json_string = {}
            # 读取excel，更新prompt
            prompt = self.prompt
            # print(template)
            # print(model_name)
            # print(input_path)
            # TODO: template合法性检查
            # 要求template为一个字典，包含field, description, key, unit, examples五个键，每个键的值为一个列表
            # 例如：
            # template = {
            #     'field': ['冷却方式', '电压'],
            #     'description': ['...', '...'],
            #     'key': [],
            #     'unit': [],
            #     'examples': []
            # }

            # 使用 pd.DataFrame.from_dict() 将字典转换为 DataFrame
            # df_template = pd.DataFrame.from_dict(template)
            #
            # # 根据 df_template 组装prompt
            # for i in range(len(df_template)):
            #     item = df_template.iloc[i]
            #     prompt += f"field: {item['field']}\n"
            #     prompt += f"\tdescription: {item['description']}\n"
            #     prompt += f"\tunit: {item['unit']}\n"
            #     prompt += f"\texample: {item['example value']}\n"
            """"""
            # 根据 self.template 组装prompt
            for i in range(len(self.template)):
                item = self.template.iloc[i]
                prompt += f"field: {item['field']}\n"
                prompt += f"\tdescription: {item['description']}\n"
                prompt += f"\tunit: {item['unit']}\n"
                prompt += f"\texample: {item['example value']}\n"

            # 处理文件路径
            # 获取文件的扩展名
            file_extension = os.path.splitext(input_path)[1].lower()
            # pdf版本
            if file_extension in ['.pdf']:
                self.download_pdf(input_path)
                # pdf_path = input_path
                table_engine = XHD_OCR_Structure()
                pdf_results = table_engine.get_text("downloaded_pdf.pdf")
                # 加上协议内容
                prompt += pdf_results
                # 更新prompt
                self.update_prompt(prompt)
            else:
                # word版本-docx
                if file_extension in ['.docx']:
                    response = requests.get(input_path)
                    if response.status_code == 200:
                        docx_content = self.download_and_read_docx(input_path)
                #         docx_content = docx_content
                #         self.update_prompt(prompt)
                #         docx_content =


            # word版本-docx
            if file_extension in ['.docx']:
                response = requests.get(input_path, stream=True)

                if response.status_code == 200:
                    with open("document.docx", "wb") as f:
                        f.write(response.content)
                else:
                    print(f"Failed to download the file. HTTP status code: {response.status_code}")

                # docx_content = self.download_and_read_docx(input_path)
                file_path = "document.docx"  # 与之前下载时保存的文件路径一致
                docx_content, tables = self.read_word_file(file_path)
                # print(docx_content)
                # 分词
                # token75s = word_tokenize(docx_content)
                #
                # # 提取前100000个tokens
                # first_100000_tokens = tokens[:50000]
                #
                # # 将tokens重新组合成文本
                # selected_text = ' '.join(first_100000_tokens)

                prompt += f"\t你需要提取的文件内容为: {docx_content}\n"
                print(prompt)
                self.update_prompt(prompt)

            #
            # # word版本-doc
            # if file_extension in ['.doc']:
            #     input_doc_path = input_path
            #     output_docx_path = "path_to_save_docx_file.docx"
            #     # self.convert_doc_to_docx(input_doc_path, output_docx_path)
            #     docx_content = self.read_docx_to_string(input_path)
            #     prompt += f"\t你需要提取的文件内容为: {docx_content}\n"
            #     self.update_prompt(prompt)
            # # todo word融合pdf

            # api
            openai.api_base = "https://openkey.cloud/v1"  # 换成代理，一定要加v1
            # openai.api_key = os.getenv("OPENAI_API_KEY")
            openai.api_key = "sk-jganxZoyZigB5qnK3725E223Ab97449198Ef9f287806E6D3"
        #
            response = openai.ChatCompletion.create(
                model = model_name,
                messages = [
                    # {"role": "system",
                    #  "content": "你是一个聪明而且有百年经验的命名实体识别（NER）识别器。"},
                    {"role": "user",
                     "content": prompt}]
            )
        #
            res = response.choices[0].message.content
            # print(res)
            # 后处理为json:去除首尾两行
            res_stripped = '\n'.join(res.splitlines()[1:-1])
            # print(res_stripped)

            # 假设res是一个JSON格式的字符串
            try:
                parsed_res = json.loads(res_stripped)
                # 如果没有抛出异常，说明res是一个有效的JSON字符串
                # 现在parsed_res是一个Python对象（如字典或列表）
                json_string = json.dumps(parsed_res)
            except json.JSONDecodeError:
                # 如果抛出异常，说明res不是一个有效的JSON字符串
                # 处理错误情况，例如设置默认值或抛出异常
                json_string = "{}"  # 创建一个空的JSON对象

        return json_string

    def read_word_file(self, file_path):
        """读取Word文档内容"""
        doc = docx.Document(file_path)

        # 获取文档中的所有文本内容
        text_content = []
        for para in doc.paragraphs:
            text_content.append(para.text)

        # 获取文档中的所有表格及其数据
        table_data = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            table_data.append(rows)

        return text_content, table_data

    def download_and_read_docx(self, url):
        try:
            # 使用 requests.get 方法下载文件
            response = requests.get(url)
            response.raise_for_status()  # 检查响应状态码

            # 指定文件保存路径
            file_path = 'qaq.docx'

            # 保存文件
            with open(file_path, 'wb') as file:
                file.write(response.content)

            # 使用 python-docx 读取 Word 文档内容
            document = Document(file_path)
            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)
            return content

        except requests.exceptions.HTTPError as e:
            return {'status': 'error', 'message': f'HTTP错误：{e}'}
        except Exception as e:
            return {'status': 'error', 'message': f'处理过程中发生异常：{e}'}

    # 调用函数示例
    # content = download_and_read_docx('您的文件URL')
    # print(content)

    def download_and_read_file(self, url):
        try:
            # 使用 requests.get 方法下载文件
            response = requests.get(url)
            response.raise_for_status()  # 检查响应状态码

            # 指定文件保存路径
            file_path = 'qaq.docx'

            # 打开文件并写入内容
            with open(file_path, 'wb') as file:
                file.write(response.content)

            # 使用 chardet 检测文件编码
            with open(file_path, 'rb') as file:
                result = chardet.detect(file.read())
                encoding = result['encoding']

            # 以检测到的编码打开文件并读取内容
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
                return content

        except requests.exceptions.HTTPError as e:
            return {'status': 'error', 'message': f'HTTP错误：{e}'}
        except Exception as e:
            return {'status': 'error', 'message': f'处理过程中发生异常：{e}'}

    def download_pdf(self, input_path, output_file_name='downloaded_pdf.pdf'):
        """
        从指定的URL下载PDF文件并保存到本地。

        :param input_path: PDF文件的URL路径。
        :param output_file_name: 保存文件的名称，默认为'downloaded_pdf.pdf'。
        """
        # 发送GET请求下载文件
        response = requests.get(input_path)

        # 检查请求是否成功
        if response.status_code == 200:
            # 打开文件并写入内容
            with open(output_file_name, 'wb') as file:
                file.write(response.content)
            print(f"文件已保存为 {output_file_name}")
        else:
            print(f"下载失败，状态码：{response.status_code}")

    # 测试函数
    # content = download_and_read_file("https://example.com/file.txt")
    # print(content)

    # def download_file_open(self, url):
    #     # print(url)
    #     # 使用 requests.get 方法下载文件
    #     response = requests.get(url)
    #     # print(type(response))
    #     # file_stream = BytesIO(response.content)
    #     # print(type(file_stream))
    #     # doc = Document(file_stream)
    #     #
    #     # for para in doc.paragraphs:
    #     #     print(para.text)
    #     # 检查响应状态码，确保下载成功
    #     if response.status_code == 200:
    #
    #         # 指定文件保存路径
    #         file_path = 'qaq.txt'
    #
    #         # 打开文件并写入内容
    #         with open(file_path, 'w') as file:
    #             file.write(response.content)
    #
    #         # 打开文件（根据文件类型，您可能需要相应的软件来打开它）
    #         # 在这里，我们只是演示如何用 Python 打开文件
    #         with open(file_path, 'r', encoding='utf-8') as file:
    #             content = file.read()
    #             return (content)
    #     else:
    #         print('下载文件失败，状态码：', response.status_code)
    #         return response.status_code

    @classmethod
    # def download_file_from_url(self, url, local_path):
    #     # # 发送GET请求并获取响应内容
    #     # response = requests.get(url)
    #     # # 检查请求是否成功（HTTP状态码为200表示成功）
    #     # if response.status_code == 200:
    #     #     # 获取文件名（从URL中提取或自定义）
    #     #     filename = url.split('/')[-1]  # 假设URL末尾就是文件名
    #     #
    #     #     # 将响应内容写入本地文件
    #     #     with open(filename, 'wb') as f:
    #     #         f.write(response.content)
    #     #
    #     #     print(f"Word文档已成功下载至{filename}")
    #     #     return filename
    #     # else:
    #     #     print(f"下载失败，HTTP状态码：{response.status_code}")
    #
    #     # 从网络URL下载文件并保存到本地路径
    #     response = requests.get(url)
    #     with open(local_path, 'wb') as f:
    #         f.write(response.content)
    def convert_doc_to_docx(input_path, output_path):
        # 启动Word应用程序
        word = client.DispatchEx("Word.Application")
        try:
            # 打开源文件
            doc = word.Documents.Open(input_path)
            # 另存为.docx
            doc.SaveAs(output_path, FileFormat=16)  # 16代表.docx格式
            doc.Close()
        finally:
            # 退出Word应用程序
            word.Quit()
    # def read_docx_to_string(self, file_path):
    #     # 加载Word文档
    #     doc = Document(file_path)
    #     full_text = []
    #
    #     # 遍历文档中的段落并读取内容
    #     for para in doc.paragraphs:
    #         full_text.append(para.text)
    #
    #     # 将所有段落连接成一个字符串
    #     return '\n'.join(full_text)


